
import re
from docx import Document
from docx.text.paragraph import Paragraph

SPHERES = {
    "Устатковання, .... маса яких становить більше одного відсотка значення нормативів порогових мас другого класу об’єктів підвищеної небезпеки":
        "Устатковання, пов’язане з виробництвом (виготовленням), використанням, переробкою, зберіганням, транспортуванням, застосуванням, утилізацією чи знешкодженням вибухопожежонебезпечних і небезпечних речовин 1 і 2 класу небезпеки, маса яких становить більше одного відсотка значення нормативів порогових мас другого класу об’єктів підвищеної небезпеки.",
    "Устатковання ... маса яких дорівнює або перевищує нормативи порогових мас":
        "Устатковання, пов’язане з виробництвом (виготовленням), використанням, переробкою, зберіганням, транспортуванням, застосуванням, утилізацією чи знешкодженням вибухопожежонебезпечних і небезпечних речовин 1 і 2 класу небезпеки, маса яких дорівнює або перевищує нормативи порогових мас.",
    "Обладнання, що працює під тиском, яке зазначене у Технічному регламенті обладнання, що працює під тиском":
        "Обладнання, що працює під тиском, яке зазначене у Технічному регламенті обладнання, що працює під тиском."
}

GROUPS_ZVIT = ["Янковський Р.А.", "Гавриленко Д.М.", "Сараєв А.Л.", "Кровяков І.В."]
GROUPS_ROZP = [
    "Начальник управління з інспектування, технічний керівник ОІ Янковський Р.А.",
    "Начальник лабораторії Гавриленко Д.М.",
    "Провідний інженер Сараєв А.Л.",
    "Провідний інженер Кровяков І.В."
]
GROUPS_ROZP_OI = ["Ірина КРАСНОПЕРОВА"]
RESPONSIBLES = ["Роман ЯНКОВСЬКИЙ", "Дмитро ГАВРИЛЕНКО"]
LEADERS = ["Янковський Р.А.", "Гавриленко Д.М."]

def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\\\/*?:"<>|]', "_", name).strip()

def clear_paragraph(paragraph: Paragraph):
    p_element = paragraph._element
    for child in list(p_element):
        p_element.remove(child)
Paragraph.clear = clear_paragraph

def replace_text_in_paragraph(paragraph, replacements):
    underline_placeholders = ["№ЗАЯВКИ", "ДАТА_ЗАЯВКИ", "ОБЄКТ", "СФЕРА"]

    full_text = "".join(run.text for run in paragraph.runs)
    full_text_norm = full_text.replace("\u00A0", " ").replace("\u202F", " ")

    underline_pattern = "|".join([rf"(?:№\s*)?{re.escape(ph)}" for ph in underline_placeholders if ph in replacements])
    if underline_pattern and re.search(underline_pattern, full_text_norm):
        paragraph.clear()
        pos = 0
        for m in re.finditer(underline_pattern, full_text_norm):
            paragraph.add_run(full_text_norm[pos:m.start()])
            ph = re.sub(r"^№\s*", "№", m.group(0)).strip()
            key = next((ph_key for ph_key in underline_placeholders if ph_key in ph), None)
            val = replacements.get(key, "")
            run = paragraph.add_run(val if val else " ")
            run.font.underline = True
            pos = m.end()
        paragraph.add_run(full_text_norm[pos:])

    for placeholder, val in replacements.items():
        ph_norm = placeholder.replace("\u00A0", " ").replace("\u202F", " ")
        pattern = rf"(?:№\s*)?{re.escape(ph_norm)}"

        if placeholder == "ВІДПОВІДАЛЬНИЙ" and re.search(pattern, full_text_norm):
            paragraph.clear()
            parts = re.split(pattern, full_text_norm)
            for i, part in enumerate(parts):
                if part:
                    paragraph.add_run(part)
                if i < len(parts) - 1:
                    paragraph.add_run("\t")
                    run_tabs_under = paragraph.add_run("\t\t")
                    run_tabs_under.font.underline = True
                    paragraph.add_run("\t\t")
                    run_val = paragraph.add_run(val if val else " ")
                    run_val.font.underline = True
            continue

        for run in paragraph.runs:
            run_text_norm = run.text.replace("\u00A0", " ").replace("\u202F", " ")
            if re.search(pattern, run_text_norm):
                run.text = re.sub(pattern, val, run_text_norm)
                run.font.underline = False

def replace_in_tables(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, replacements)

def replace_placeholders(doc_path: str, replacements: dict, output_path: str):
    doc = Document(doc_path)
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, replacements)
    replace_in_tables(doc, replacements)
    doc.save(output_path)
